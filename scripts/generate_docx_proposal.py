import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    # top, bottom, left, right in dxa (1/20 of a pt)
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m_type, m_val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m_type}')
        node.set(qn('w:w'), str(m_val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_left_border(cell, color_hex="1B365D", sz="36"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    
    # Left border
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)  # size: 36 = 4.5 pt
    left.set(qn('w:space'), '0')
    left.set(qn('w:color'), color_hex)
    borders.append(left)
    
    # Clear other borders
    for b_type in ['top', 'bottom', 'right']:
        node = OxmlElement(f'w:{b_type}')
        node.set(qn('w:val'), 'none')
        borders.append(node)
        
    tcPr.append(borders)

def make_callout_box(doc, text, title=""):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    cell = table.cell(0, 0)
    cell.width = Inches(6.0)
    
    set_cell_background(cell, "F4F6F8")
    set_cell_left_border(cell, color_hex="1B365D", sz="36")
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    
    if title:
        run_title = p.add_run(title + "\n")
        run_title.font.name = 'Segoe UI'
        run_title.font.size = Pt(11)
        run_title.font.bold = True
        run_title.font.color.rgb = RGBColor(27, 54, 93) # #1B365D
        
    run_text = p.add_run(text)
    run_text.font.name = 'Segoe UI'
    run_text.font.size = Pt(10.5)
    run_text.font.italic = True
    run_text.font.color.rgb = RGBColor(80, 80, 80)
    
    # Add empty paragraph after table for spacing
    p_after = doc.add_paragraph()
    p_after.paragraph_format.space_before = Pt(0)
    p_after.paragraph_format.space_after = Pt(6)

def add_styled_heading(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(27, 54, 93) # #1B365D
        
        # Add a subtle bottom border element using XML if possible, or just text representation
        # We can draw a horizontal divider line below H1
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')  # 0.75 pt
        bottom.set(qn('w:space'), '4')
        bottom.set(qn('w:color'), 'A0ABBA')  # Muted blue-grey
        pBdr.append(bottom)
        p._p.get_or_add_pPr().append(pBdr)
        
    elif level == 2:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 90, 156) # #005A9C
        
    elif level == 3:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        run.font.name = 'Segoe UI'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = RGBColor(80, 80, 80)
        
    return p

def add_body_paragraph(doc, text, bold_prefix="", space_after=6, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.font.name = 'Segoe UI'
        r_prefix.font.size = Pt(11)
        r_prefix.font.bold = True
        r_prefix.font.color.rgb = RGBColor(40, 40, 40)
        
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
    run.font.italic = italic
    run.font.color.rgb = RGBColor(50, 50, 50)
    return p

def add_bullet_point(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    
    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        r_prefix.font.name = 'Segoe UI'
        r_prefix.font.size = Pt(11)
        r_prefix.font.bold = True
        r_prefix.font.color.rgb = RGBColor(40, 40, 40)
        
    run = p.add_run(text)
    run.font.name = 'Segoe UI'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(50, 50, 50)
    return p

def main():
    doc = Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Document Header / Cover style
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(2)
    title_run = title_p.add_run("COGNITIVE SAFETY & RUNTIME GOVERNANCE")
    title_run.font.name = 'Segoe UI'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(27, 54, 93) # #1B365D
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.paragraph_format.space_before = Pt(0)
    subtitle_p.paragraph_format.space_after = Pt(12)
    sub_run = subtitle_p.add_run("A Joint Safety Framework Utilizing Błyskawica V8 and Nethical")
    sub_run.font.name = 'Segoe UI'
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0, 90, 156) # #005A9C
    
    proposal_p = doc.add_paragraph()
    proposal_p.paragraph_format.space_after = Pt(24)
    prop_run = proposal_p.add_run("A Collaborative Safety & Security Proposal for the UK Artificial Intelligence Security Institute (UK AISI)")
    prop_run.font.name = 'Segoe UI'
    prop_run.font.size = Pt(11)
    prop_run.font.italic = True
    prop_run.font.bold = True
    prop_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # Contact Info Table
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    info_table.autofit = False
    
    labels = [
        "Architect / Developer:", 
        "Primary Location:", 
        "Direct Contact:", 
        "Email Address:", 
        "GitHub Portfolio:"
    ]
    values = [
        "Andrzej Matewski (aka V1B3hR), Freelance Systems Architect & Engineer",
        "Preston, Lancashire, United Kingdom",
        "07912853241 (International: +44 7912 853 241)",
        "brightnightbeacon@gmail.com",
        "https://github.com/V1B3hR (Key Framework: https://github.com/V1B3hR/nethical)"
    ]
    
    for i in range(5):
        row = info_table.rows[i]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        
        cell_lbl.width = Inches(1.8)
        cell_val.width = Inches(4.5)
        
        # Style label
        lbl_p = cell_lbl.paragraphs[0]
        lbl_p.paragraph_format.space_after = Pt(2)
        lbl_run = lbl_p.add_run(labels[i])
        lbl_run.font.name = 'Segoe UI'
        lbl_run.font.size = Pt(9.5)
        lbl_run.font.bold = True
        lbl_run.font.color.rgb = RGBColor(80, 80, 80)
        
        # Style value
        val_p = cell_val.paragraphs[0]
        val_p.paragraph_format.space_after = Pt(2)
        val_run = val_p.add_run(values[i])
        val_run.font.name = 'Segoe UI'
        val_run.font.size = Pt(9.5)
        val_run.font.color.rgb = RGBColor(50, 50, 50)
        if "https://" in values[i]:
            val_run.font.underline = True
            val_run.font.color.rgb = RGBColor(0, 90, 156)
            
        set_cell_margins(cell_lbl, top=40, bottom=40, left=60, right=60)
        set_cell_margins(cell_val, top=40, bottom=40, left=60, right=60)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Executive Summary
    add_styled_heading(doc, "Executive Summary", level=1)
    
    add_body_paragraph(doc, 
        "As artificial intelligence transitions from passive query-response interfaces to autonomous agentic architectures operating in real-world environments, the paradigms for AI safety, compliance, and runtime alignment must evolve accordingly. Classic post-hoc safety guardrails (such as wrapper filters and semantic classifiers) are prone to jailbreaking, context drift, and adversarial attacks.", 
        space_after=6
    )
    
    add_body_paragraph(doc, 
        "This proposal outlines a dual-layered, physically grounded AI safety and alignment ecosystem designed to address these concerns at their root:",
        space_after=6
    )
    
    add_bullet_point(doc, 
        "A hybrid cognitive modeling and neuromorphic simulation network. Rather than merely making statistical 'guesses' based on language correlations, Błyskawica anchors its core learning representations in physical hardware entropy (IBM Quantum) and strictly enforces the invariant laws of physics (PINN Thermodynamics) in its training passes. It mathematically prevents generative hallucinations through Euclidean-bounded thought networks.",
        bold_prefix="1. Błyskawica V8 (Lightning): "
    )
    
    add_bullet_point(doc, 
        "An open-source, local-first runtime governance and threat detection gateway. Positioned between autonomous agents and their operating environment, Nethical intercepts proposed agent actions in real time, screens them against the 25 Fundamental Laws of AI, and produces cryptographically verifiable, append-only audit proofs with sub-50ms latency.",
        bold_prefix="2. Nethical Governance Layer: "
    )
    
    add_body_paragraph(doc, 
        "By merging Błyskawica's physical-thermodynamic constraints with Nethical's deterministic runtime enforcement, we establish a robust framework for high-stakes environments. We submit this technical brief to the UK Artificial Intelligence Security Institute (UK AISI) to explore collaborative avenues in empirical safety benchmarking, quantum-grounded alignment validation, and the development of local-first auditing standards.",
        space_after=12
    )
    
    # About the Architect
    add_styled_heading(doc, "About the Architect: Practical Engineering-First AI Design", level=1)
    
    add_body_paragraph(doc, 
        "Andrzej Matewski (known globally in the open-source and developer community as V1B3hR) is an independent systems architect and software engineer. Educated as an electrical technician with a mid-advanced academic foundation (Matura), his approach to artificial intelligence is fundamentally practical, hardware-grounded, and systems-focused, departing from purely abstract statistical theories.",
        space_after=6
    )
    
    add_body_paragraph(doc, 
        "Drawing direct analogies from power transmission, feedback loops, electrical compensating chokes, and AC filter circuits, Matewski designs AI systems where stability, safety, and energy constraints are treated as physical constants. This methodology translates into several software engineering qualities:",
        space_after=6
    )
    
    add_bullet_point(doc, 
        "The codebase is built for execution in high-stakes environments. It features 100% test coverage across specialized cognitive intelligence, biological plausibility, and emergent behavior suites (43+ formal tests), containerized OpenShift environments, and hardware acceleration via AMD ROCm Instinct GPU Matrix Cores.",
        bold_prefix="Production-Ready Rigor: "
    )
    
    add_bullet_point(doc, 
        "The model is bound by physical equations (e.g., 1D Heat Conduction) to make predictions mathematically consistent with thermodynamics, minimizing energy waste and component wear in industrial deployments.",
        bold_prefix="Physical Grounding: "
    )
    
    add_bullet_point(doc, 
        "All code, tools, and regulatory guidelines (such as the 25 Fundamental Laws) are open-source, promoting global transparency, peer review, and public reproducibility.",
        bold_prefix="Open-Source Commitment: "
    )
    
    # Section 1: The Błyskawica Cognitive Architecture
    add_styled_heading(doc, "Section 1: The Błyskawica Cognitive Architecture (Lightning)", level=1)
    
    add_body_paragraph(doc, 
        "Błyskawica V8 is a complex cognitive-simulation architecture rather than a simple mathematical statistical model. It represents a paradigm shift by coupling artificial neural parameters with physical constants and simulated biological mechanisms to ensure cognitive stability and absolute safety.",
        space_after=8
    )
    
    make_callout_box(doc, 
        "In classical approaches, AI merely 'guesses' answers based on statistical distributions. Błyskawica V8 breaks away from this by anchoring its cognitive layers directly in physical quantum hardware (IBM Quantum) and the invariant laws of physics (PINN Thermodynamics), resulting in a highly secure, hallucination-resistant system.",
        title="Błyskawica Architectural Core Directive"
    )
    
    add_styled_heading(doc, "Key Scientific Innovations and Operational Value:", level=2)
    
    add_bullet_point(doc, 
        "By utilizing quantum entanglement and superposition (Hadamard and CNOT gates) on physical IBM Quantum processors, Błyskawica extracts true physical entropy (quantum noise) to set the foundational weights of its neural layers, eliminating statistical bias and preventing mathematical lock-ins.",
        bold_prefix="1. Quantum Baptism (Hardware Entropy Integration): "
    )
    
    add_bullet_point(doc, 
        "Integrates Fourier's 1D Heat Conduction Equation (u_t - α u_xx = 0) directly into the backpropagation pass of a PyTorch training engine. The model predicts thermal wear and remaining useful life (validated on industrial datasets) not via heuristic curve-fitting, but by strictly adhering to the conservation of thermal energy. This decreases physical prediction loss by 47% in 10 epochs, rendering it ideal for high-precision Industrial Digital Twins.",
        bold_prefix="2. Physics-Informed Neural Networks (PINN Thermodynamics): "
    )
    
    add_bullet_point(doc, 
        "Implements a double-loop feedback cycle: a low-latency 'Reflex' loop for instinctive, instantaneous context parsing, and a slow, deep 'Reflection' loop for ethical and logical reasoning, stabilized by digital filters modeled on AC compensation chokes. This optimizes background memory footprint to ~4.8GB of VRAM.",
        bold_prefix="3. Dual-Rotor Cognitive Engine: "
    )
    
    add_bullet_point(doc, 
        "Establishes a recursive tree structure for thought expansion where every branch is constantly monitored for its Euclidean distance from the factual root. If the distance exceeds 50.0 units, the RealityAnchor safety circuit halts recursion and grounds the model back to verified facts, mathematically eliminating generative hallucinations (AI psychosis).",
        bold_prefix="4. Reality Anchor & Fractal Garden: "
    )
    
    add_bullet_point(doc, 
        "An asynchronous memory watchdog that serializes virtual neurochemical states, relational parameters, and active context paths in milliseconds. This guarantees zero operational downtime and complete recovery resilience against sudden power losses or hardware failures.",
        bold_prefix="5. Memory Shield (MemoryGuard): "
    )
    
    add_bullet_point(doc, 
        "Maps 1D EEG inputs (8–12 Hz Alpha waves) to 2D Chladni resonant grids on a 16x16 Diamond Yant matrix. Under focus, this coherence dynamically suppresses physical IBM Phase Change Memory (PCM) resistance drift by 49.2%, stabilizing memristive hardware storage via human-in-the-loop biological feedback.",
        bold_prefix="6. Diamond Yant & Cognitive Cymatics: "
    )
    
    add_bullet_point(doc, 
        "Simulates w-system neurochemical levels (Dopamine, Serotonin, Melatonin, Cortisol) to modulate learning speeds and stress thresholds. It autonomously scales down training rates during simulated rest phases to prevent overfitting, and boosts focus during active computations.",
        bold_prefix="7. Autonomous Neurochemical Self-Regulation: "
    )
    
    add_bullet_point(doc, 
        "Integrates dipole-dipole quantum coupling models within tubulin dimers (Orch OR quantum biology). Proves that focused neurochemical parameters damp thermal phonon noise, extending simulated quantum coherence survival times by 74.5% (from 10.25 ps to 17.89 ps) in warm, biological environments.",
        bold_prefix="8. Microtubule Quantum Coherence: "
    )
    
    # Section 2: The Nethical Governance Layer
    add_styled_heading(doc, "Section 2: The Nethical Governance & Ethics Layer", level=1)
    
    add_body_paragraph(doc, 
        "Nethical (https://github.com/V1B3hR/nethical) is an open-source, local-first AI governance and runtime safety framework. It is designed to act as an intercepting gateway between autonomous agents (LLMs, assistants, or robotic systems) and the external world, ensuring that every action is audited and compliant before execution.",
        space_after=8
    )
    
    add_styled_heading(doc, "Core Architectural Principles:", level=2)
    
    add_bullet_point(doc, 
        "An immutable, built-in digital 'Bill of Rights and Duties' for AI. Every agent request is evaluated against these laws at runtime, returning an immediate decision (ALLOW, RESTRICT, BLOCK, TERMINATE). Upstream modifications to these laws require a formal two-reviewer consensus to prevent value drift.",
        bold_prefix="1. The 25 Fundamental Laws: "
    )
    
    add_bullet_point(doc, 
        "Safety must not fail when network connectivity is lost. Nethical is designed to run locally on the host machine or edge device, enforcing ethics and security policies without external API dependencies.",
        bold_prefix="2. Local-First Design: "
    )
    
    add_bullet_point(doc, 
        "Every evaluation returns a 'Decision + Reason + Proof' tuple. The proof is cryptographically anchored in an append-only, Merkle-tree-structured audit log, creating a tamper-evident record of all AI behavior for regulatory compliance.",
        bold_prefix="3. Tamper-Evident Auditing: "
    )
    
    add_styled_heading(doc, "Real-Time Threat Detection Suite (Sub-50ms Latency):", level=2)
    
    add_body_paragraph(doc, 
        "To ensure that safety enforcement does not bottleneck performance, Nethical integrates five specialized realtime detectors optimized for ultra-low latency execution:",
        space_after=6
    )
    
    add_bullet_point(doc, 
        "Scans infrastructure for unauthorized model execution (Ollama, LM Studio, vLLM) or API calls (OpenAI, Anthropic, Google) to prevent corporate data leakage.",
        bold_prefix="• Shadow AI Detector (<20ms): "
    )
    
    add_bullet_point(doc, 
        "Applies multi-modal frequency and temporal analysis to detect face swaps, GAN artifacts, and voice cloning in incoming or outgoing media streams.",
        bold_prefix="• Deepfake Detector (<30ms): "
    )
    
    add_bullet_point(doc, 
        "Analyzes system call sequences and memory access patterns to block mutating, polymorphic exploits launched by compromised agents.",
        bold_prefix="• Polymorphic Malware Detector (<50ms): "
    )
    
    add_bullet_point(doc, 
        "Protects agents from direct jailbreaks (DAN, Apophis), indirect injection vectors, and system prompt leakage using a fast two-tier validation mechanism.",
        bold_prefix="• Prompt Injection Guard (<15ms): "
    )
    
    add_bullet_point(doc, 
        "Defends against model extraction attempts, membership inference, and adversarial inputs designed to bypass neural classifiers.",
        bold_prefix="• AI vs AI Defender (<25ms): "
    )
    
    # Section 3: Proposed Collaboration with UK AISI
    add_styled_heading(doc, "Section 3: Proposed Areas of Collaboration with UK AISI", level=1)
    
    add_body_paragraph(doc, 
        "The UK Artificial Intelligence Security Institute (UK AISI) represents the vanguard of scientific AI safety. We propose establishing a collaborative partnership to advance the science of physical alignment and local-first runtime safety in several key areas:",
        space_after=6
    )
    
    add_bullet_point(doc, 
        "Utilizing UK AISI's advanced red-teaming methodologies to stress-test Błyskawica's Reality Anchor and Nethical's Prompt Injection Guard against state-of-the-art attack models.",
        bold_prefix="1. Empirical Safety Benchmarking: "
    )
    
    add_bullet_point(doc, 
        "Investigating the mathematical and security benefits of quantum-entangled entropy weights as an alignment mechanism. We propose analyzing whether grounding model states in physical quantum entropy prevents the 'creative exploitation' of safety boundaries seen in classical networks.",
        bold_prefix="2. Quantum-Grounded Alignment Research: "
    )
    
    add_bullet_point(doc, 
        "Co-developing open-source technical specifications for local-first AI gateways, using Nethical's 'Decision + Reason + Proof' Merkle log as a reference standard for auditable autonomous agents.",
        bold_prefix="3. Open Standards for Runtime Auditing: "
    )
    
    add_bullet_point(doc, 
        "Testing the deployment of PINN Thermodynamics and Nethical safety gateways in cyber-physical systems and industrial digital twins where safety-critical decisions cannot tolerate cloud latency or network instability.",
        bold_prefix="4. Safety-Critical Edge Deployments: "
    )
    
    doc.save("Blyskawica_Nethical_Proposal.docx")
    print("SUCCESS: Generated Blyskawica_Nethical_Proposal.docx")

if __name__ == "__main__":
    main()
